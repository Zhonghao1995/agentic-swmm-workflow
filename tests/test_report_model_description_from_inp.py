"""Regression tests: the report's Model Description comes from the run's INP.

Bug (found 2026-08-08 manifest-drift sweep, HIGH): the default report
template's Model Description section read ``basin_area_ha`` /
``sim_start`` / ``sim_end`` / ``landuse_params`` / ``green_ampt_params``
from the top ``manifest.json``, a schema no production writer emits
(only the one-off todcreek benchmark script hand-writes it). Every run
from either real pipeline rendered ``n/a`` throughout the table while
the narrative claimed the values "reflect the exact inputs used".

Fix under test: facts are parsed from the run's own INP (total
subcatchment area, area-weighted imperviousness, [OPTIONS] simulation
window, Green-Ampt averages gated on the INFILTRATION method), with
the legacy manifest keys kept as a fallback so hand-written benchmark
manifests keep rendering. The narrative states the real provenance.
"""

from __future__ import annotations

import json
import subprocess
import sys
import unittest
from pathlib import Path
from tempfile import TemporaryDirectory

from docx import Document


_REPO = Path(__file__).resolve().parents[1]
_SCRIPT = str(_REPO / "skills/swmm-report/scripts/generate_report.py")


_INP = """\
[OPTIONS]
FLOW_UNITS           LPS
INFILTRATION         GREEN_AMPT
START_DATE           01/01/1994
START_TIME           00:00:00
END_DATE             01/03/1994
END_TIME             12:00:00

[SUBCATCHMENTS]
;;Name  RainGage  Outlet  Area  %Imperv  Width  Slope  CurbLen
S1      RG1       J1      10.0  80.0     100    0.5    0
S2      RG1       J2      30.0  20.0     200    0.5    0

[INFILTRATION]
;;Subcatch  Suction  Ksat  IMD
S1          90.0     10.0  0.25
S2          110.0    6.0   0.25
"""


def _make_run_dir(tmp: Path) -> str:
    run_dir = tmp / "run"
    (run_dir / "09_audit").mkdir(parents=True)
    (run_dir / "05_builder").mkdir()
    (run_dir / "05_builder" / "model.inp").write_text(_INP, encoding="utf-8")
    # Realistic production top manifest: NO fictional model keys.
    (run_dir / "manifest.json").write_text(
        json.dumps({"schema_version": "1.0", "generated_by": "aiswmm-run", "outputs": {}}),
        encoding="utf-8",
    )
    (run_dir / "09_audit" / "experiment_provenance.json").write_text(
        json.dumps(
            {
                "run_id": "inp-facts",
                "status": "pass",
                "qa": {"status": "pass", "fail_count": 0, "pass_count": 1, "checks": []},
                "metrics": {
                    "peak_flow": {"value": 2.5, "time_hhmm": "03:15"},
                    "continuity_error": -0.2,
                    "swmm_return_code": 0,
                },
                "artifacts": {},
            }
        ),
        encoding="utf-8",
    )
    (run_dir / "model_diagnostics.json").write_text(
        json.dumps({"schema_version": "1.0", "status": "pass", "error_count": 0,
                    "warning_count": 0, "diagnostics": []}),
        encoding="utf-8",
    )
    (run_dir / "09_audit" / "comparison.json").write_text(
        json.dumps({"comparison_available": False, "checks": [], "warnings": []}),
        encoding="utf-8",
    )
    return str(run_dir)


def _all_text(doc: Document) -> str:
    texts = []
    for para in doc.paragraphs:
        if para.text.strip():
            texts.append(para.text.strip())
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    texts.append(cell.text.strip())
    return "\n".join(texts)


class ModelDescriptionFromInpTests(unittest.TestCase):
    def _generate(self, tmp: Path) -> str:
        run_dir = _make_run_dir(tmp)
        out = str(tmp / "report.docx")
        result = subprocess.run(
            [sys.executable, _SCRIPT, "--run-dir", run_dir, "--out", out],
            capture_output=True,
            text=True,
        )
        self.assertEqual(result.returncode, 0, result.stderr)
        return _all_text(Document(out))

    def test_model_description_values_derive_from_inp(self) -> None:
        """Pre-fix: every one of these cells rendered n/a."""
        with TemporaryDirectory() as tmp:
            text = self._generate(Path(tmp))
        # Total area 40.0 ha; area-weighted imperviousness
        # (10*80 + 30*20) / 40 = 35.0 %.
        self.assertIn("40.0", text)
        self.assertIn("35.0", text)
        # Simulation window from [OPTIONS].
        self.assertIn("01/01/1994 00:00:00", text)
        self.assertIn("01/03/1994 12:00:00", text)
        # Green-Ampt averages: suction (90+110)/2, ksat (10+6)/2.
        self.assertIn("100.0", text)
        self.assertIn("8.0", text)

    def test_narrative_no_longer_claims_manifest_exactness(self) -> None:
        with TemporaryDirectory() as tmp:
            text = self._generate(Path(tmp))
        self.assertNotIn("exact inputs used", text)
        self.assertIn("parsed from the run's own SWMM INP", text)

    def test_horton_model_reports_na_for_green_ampt_fields(self) -> None:
        """Under Horton the INFILTRATION columns mean something else and
        must not be reported as suction/ksat."""
        with TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            run_dir = _make_run_dir(tmp_path)
            inp_path = Path(run_dir) / "05_builder" / "model.inp"
            inp_path.write_text(
                _INP.replace("GREEN_AMPT", "HORTON"), encoding="utf-8"
            )
            out = str(tmp_path / "report.docx")
            result = subprocess.run(
                [sys.executable, _SCRIPT, "--run-dir", run_dir, "--out", out],
                capture_output=True,
                text=True,
            )
            self.assertEqual(result.returncode, 0, result.stderr)
            text = _all_text(Document(out))
        # Area/window still derive; infiltration cells fall to n/a.
        self.assertIn("40.0", text)
        self.assertIn("n/a", text)

    def test_legacy_handwritten_manifest_still_renders(self) -> None:
        """Benchmark runs that hand-write the legacy keys keep working
        when no INP is present in the run dir."""
        with TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            run_dir = _make_run_dir(tmp_path)
            (Path(run_dir) / "05_builder" / "model.inp").unlink()
            (Path(run_dir) / "manifest.json").write_text(
                json.dumps(
                    {
                        "basin_area_ha": 1858.75,
                        "sim_start": "1984-05-23",
                        "sim_end": "1984-05-27",
                        "landuse_params": {"imperv_pct": 25.24},
                        "green_ampt_params": {
                            "suction_mm": 90.82,
                            "ksat_mm_per_hr": 8.9,
                        },
                    }
                ),
                encoding="utf-8",
            )
            out = str(tmp_path / "report.docx")
            result = subprocess.run(
                [sys.executable, _SCRIPT, "--run-dir", run_dir, "--out", out],
                capture_output=True,
                text=True,
            )
            self.assertEqual(result.returncode, 0, result.stderr)
            text = _all_text(Document(out))
        self.assertIn("1858.75", text)
        self.assertIn("25.24", text)


if __name__ == "__main__":  # pragma: no cover
    unittest.main()
