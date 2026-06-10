"""Tests for skills/swmm-report/scripts/generate_report.py.

All tests use a synthetic fixture run dir constructed in tmp_path.
No dependency on real run dirs under runs/ — tests are self-contained.

python-docx is required for the readback assertions; if not installed the
entire module is skipped via pytest.importorskip.
"""

import json
import os
import re
import struct
import subprocess
import sys
import textwrap
import zlib

import pytest

docx = pytest.importorskip("docx")
yaml = pytest.importorskip("yaml")

from docx import Document  # noqa: E402  (after importorskip guard)
from docx.oxml.ns import qn  # noqa: E402

SCRIPT = os.path.join(
    os.path.dirname(__file__),
    "..",
    "skills",
    "swmm-report",
    "scripts",
    "generate_report.py",
)
DEFAULT_TEMPLATE = os.path.join(
    os.path.dirname(__file__),
    "..",
    "skills",
    "swmm-report",
    "templates",
    "default.yaml",
)


# ---------------------------------------------------------------------------
# Minimal 1-pixel PNG (inline — no PIL dependency)
# ---------------------------------------------------------------------------

def _make_1px_png() -> bytes:
    """Return the bytes of a valid 1x1 white RGB PNG."""
    def chunk(name: bytes, data: bytes) -> bytes:
        c = struct.pack(">I", len(data)) + name + data
        return c + struct.pack(">I", zlib.crc32(name + data) & 0xFFFFFFFF)

    signature = b"\x89PNG\r\n\x1a\n"
    ihdr_data = struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0)
    ihdr = chunk(b"IHDR", ihdr_data)
    raw_row = b"\x00\xff\xff\xff"  # filter byte 0 + RGB white
    compressed = zlib.compress(raw_row)
    idat = chunk(b"IDAT", compressed)
    iend = chunk(b"IEND", b"")
    return signature + ihdr + idat + iend


# ---------------------------------------------------------------------------
# Fixture factory
# ---------------------------------------------------------------------------

MINIMAL_PROVENANCE = {
    "schema_version": "1.0",
    "run_id": "test-run-01",
    "generated_at_utc": "2026-01-15T08:00:00+00:00",
    "generated_by": "swmm-experiment-audit",
    "case_name": "test-run-01",
    "status": "pass",
    "metrics": {
        "peak_flow": {
            "name": "peak_flow",
            "node": "O1",
            "value": 1.184,
            "unit": "CMS",
            "time_hhmm": "12:47",
            "source_artifact": "runner_rpt",
            "source_field": "Maximum Total Inflow",
            "source_section": "Node Inflow Summary",
            "source_validation": {"matches_report": True},
        },
        "continuity_error": None,
        "swmm_return_code": 0,
        "builder_counts": None,
    },
    "qa": {
        "status": "pass",
        "pass_count": 2,
        "fail_count": 0,
        "checks": [
            {"id": "runner_outputs_exist", "ok": True, "detail": "rpt_exists=True out_exists=True"},
            {"id": "peak_metric_present", "ok": True, "detail": "source_section=Node Inflow Summary"},
        ],
    },
    "artifacts": {
        "model_inp": {
            "id": "model_inp",
            "exists": True,
            "role": "SWMM input model",
            "sha256": "4840dbe4d20744889fecdc075031187af87b431cce86682a6be8f56801f81a84",
            "produced_by": "swmm-builder",
            "absolute_path": "/fake/model.inp",
            "relative_path": "runs/test/model.inp",
            "used_for": ["SWMM execution input"],
        },
    },
    "model_diagnostics": {
        "schema_version": "1.0",
        "status": "pass",
        "error_count": 0,
        "warning_count": 0,
        "diagnostics": [],
        "generated_at_utc": "2026-01-15T08:00:00+00:00",
        "generated_by": "swmm-experiment-audit",
        "source_inp": "runs/test/model.inp",
        "source_rpt": "runs/test/model.rpt",
    },
    "repo": {
        "git_head": "abc123def456",
        "git_branch": "main",
        "git_status_porcelain": "",
        "root": "/fake/repo",
    },
    "tools": {
        "python_version": "3.11.0",
        "python_executable": "/usr/bin/python3",
        "swmm5_version": "5.2.4",
    },
    "run_dir": {"absolute_path": "/fake/run", "relative_path": "runs/test"},
    "raw_sources": {},
    "inputs": {},
    "commands": [],
    "objective": None,
    "warnings": [],
    "workflow_mode": None,
    "uncertainty_ensemble": None,
}

MINIMAL_MANIFEST = {
    "run_dir": "/fake/run",
    "sim_start": "1984-05-23",
    "sim_end": "1984-05-27",
    "basin_area_ha": 1858.75,
    "landuse_params": {
        "imperv_pct": 25.24,
        "n_imperv": 0.015,
        "n_perv": 0.297,
        "dstore_imperv_mm": 1.27,
        "dstore_perv_mm": 3.02,
        "zero_imperv_pct": 11.06,
    },
    "green_ampt_params": {
        "suction_mm": 90.82,
        "ksat_mm_per_hr": 8.90,
        "imdmax": 0.251,
    },
    "qoi": {"peak_flow_cms_at_O1": 1.184, "time_of_peak_hhmm": "12:47"},
}

MINIMAL_DIAGNOSTICS = {
    "schema_version": "1.0",
    "status": "pass",
    "error_count": 0,
    "warning_count": 0,
    "diagnostics": [],
    "generated_at_utc": "2026-01-15T08:00:00+00:00",
    "generated_by": "swmm-experiment-audit",
    "source_inp": "runs/test/model.inp",
    "source_rpt": "runs/test/model.rpt",
}

COMPARISON_UNAVAILABLE = {
    "comparison_available": False,
    "current_run_id": "test-run-01",
    "generated_at_utc": "2026-01-15T08:00:00+00:00",
    "generated_by": "swmm-experiment-audit",
    "reason": "No --compare-to run directory was provided.",
    "schema_version": "1.0",
}

COMPARISON_AVAILABLE = {
    "comparison_available": True,
    "current_run_id": "test-run-01",
    "baseline_run_id": "baseline-run",
    "generated_at_utc": "2026-01-15T08:00:00+00:00",
    "generated_by": "swmm-experiment-audit",
    "schema_version": "1.0",
    "metrics": {
        "peak_flow": {"current": 1.184, "baseline": 1.050, "delta": 0.134},
    },
}


def _make_run_dir(
    tmp_path,
    *,
    with_figures: bool = False,
    comparison: dict | None = None,
    manifest: dict | None = None,
    provenance: dict | None = None,
):
    """Build a synthetic run dir under tmp_path."""
    run_dir = tmp_path / "run"
    run_dir.mkdir()
    audit_dir = run_dir / "09_audit"
    audit_dir.mkdir()

    # manifest.json
    (run_dir / "manifest.json").write_text(
        json.dumps(manifest or MINIMAL_MANIFEST), encoding="utf-8"
    )
    # 09_audit/experiment_provenance.json
    (audit_dir / "experiment_provenance.json").write_text(
        json.dumps(provenance or MINIMAL_PROVENANCE), encoding="utf-8"
    )
    # model_diagnostics.json (top-level)
    (run_dir / "model_diagnostics.json").write_text(
        json.dumps(MINIMAL_DIAGNOSTICS), encoding="utf-8"
    )
    # 09_audit/comparison.json
    comp_data = comparison if comparison is not None else COMPARISON_UNAVAILABLE
    (audit_dir / "comparison.json").write_text(
        json.dumps(comp_data), encoding="utf-8"
    )

    if with_figures:
        plot_dir = run_dir / "07_plot"
        plot_dir.mkdir()
        (plot_dir / "outfall_flow.png").write_bytes(_make_1px_png())
        (plot_dir / "junction_flow.png").write_bytes(_make_1px_png())

    return str(run_dir)


def _run_script(run_dir: str, out: str, extra_args: list[str] | None = None) -> subprocess.CompletedProcess:
    cmd = [sys.executable, SCRIPT, "--run-dir", run_dir, "--out", out]
    if extra_args:
        cmd.extend(extra_args)
    return subprocess.run(cmd, capture_output=True, text=True)


def _all_text(doc: Document) -> list[str]:
    """Extract all non-empty paragraph and table-cell texts from document."""
    texts = []
    for para in doc.paragraphs:
        if para.text.strip():
            texts.append(para.text.strip())
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    texts.append(cell.text.strip())
    return texts


def _heading_texts(doc: Document) -> list[str]:
    return [p.text.strip() for p in doc.paragraphs if p.style.name.startswith("Heading")]


def _paragraph_texts(doc: Document) -> list[str]:
    """All paragraph texts in document order (including empty)."""
    return [p.text for p in doc.paragraphs]


def _body_elements(doc: Document):
    """Yield (type, obj) for each top-level body element in order.

    type is 'para' or 'table'.
    """
    from docx.oxml.ns import qn as _qn
    body = doc.element.body
    para_tag = _qn("w:p")
    tbl_tag = _qn("w:tbl")
    for child in body:
        if child.tag == para_tag:
            from docx.text.paragraph import Paragraph
            yield "para", Paragraph(child, doc)
        elif child.tag == tbl_tag:
            from docx.table import Table
            yield "table", Table(child, doc)


# ---------------------------------------------------------------------------
# Tests
# ---------------------------------------------------------------------------

class TestHeadingsPresent:
    """Default template sections appear in document in order with numbered headings."""

    def test_all_section_headings_present(self, tmp_path):
        run_dir = _make_run_dir(tmp_path, with_figures=True, comparison=COMPARISON_UNAVAILABLE)
        out = str(tmp_path / "report.docx")
        result = _run_script(run_dir, out)
        assert result.returncode == 0, result.stderr

        doc = Document(out)
        headings = _heading_texts(doc)

        # Cover is unnumbered (Heading 1); body sections get sequential numbers
        expected = [
            "Run Audit Deliverable",            # cover (Heading 1, unnumbered)
            "1 Run Summary",
            "2 Model Description",
            "3 QA Gates",
            "4 Figures",
            "5 Model Diagnostics",
            # comparison skipped because comparison_available=false
            "6 Artifact Provenance",
            "7 Appendix: Generation Environment",
        ]
        for heading in expected:
            assert heading in headings, f"Heading {heading!r} not found in {headings}"

    def test_at_least_one_table(self, tmp_path):
        run_dir = _make_run_dir(tmp_path)
        out = str(tmp_path / "report.docx")
        result = _run_script(run_dir, out)
        assert result.returncode == 0, result.stderr
        doc = Document(out)
        assert len(doc.tables) >= 1


class TestMetricsValues:
    """Metrics table values match the manifest / provenance values."""

    def test_peak_flow_value_present(self, tmp_path):
        run_dir = _make_run_dir(tmp_path)
        out = str(tmp_path / "report.docx")
        result = _run_script(run_dir, out)
        assert result.returncode == 0, result.stderr
        doc = Document(out)
        all_text = _all_text(doc)
        assert "1.184" in all_text, f"Peak flow 1.184 not found in document text: {all_text[:30]}"

    def test_provenance_sha256_prefix_present(self, tmp_path):
        run_dir = _make_run_dir(tmp_path)
        out = str(tmp_path / "report.docx")
        result = _run_script(run_dir, out)
        assert result.returncode == 0, result.stderr
        doc = Document(out)
        all_text = " ".join(_all_text(doc))
        assert "4840dbe4" in all_text, "SHA-256 prefix 4840dbe4 not found in provenance table"


class TestFigures:
    """Figure embedding: present with PNGs, graceful note without."""

    def test_figures_embedded_when_plot_dir_exists(self, tmp_path):
        run_dir = _make_run_dir(tmp_path, with_figures=True)
        out = str(tmp_path / "report.docx")
        result = _run_script(run_dir, out)
        assert result.returncode == 0, result.stderr
        doc = Document(out)
        # At least two figure captions expected (outfall_flow, junction_flow)
        all_text = " ".join(_all_text(doc))
        assert "outfall_flow" in all_text
        assert "junction_flow" in all_text

    def test_no_figures_note_when_no_plot_dir(self, tmp_path):
        run_dir = _make_run_dir(tmp_path, with_figures=False)
        out = str(tmp_path / "report.docx")
        result = _run_script(run_dir, out)
        assert result.returncode == 0, result.stderr
        doc = Document(out)
        all_text = " ".join(_all_text(doc))
        assert "No figures available" in all_text


class TestComparisonConditional:
    """Comparison section present only when comparison_available=true."""

    def test_comparison_absent_when_unavailable(self, tmp_path):
        run_dir = _make_run_dir(tmp_path, comparison=COMPARISON_UNAVAILABLE)
        out = str(tmp_path / "report.docx")
        _run_script(run_dir, out)
        doc = Document(out)
        headings = _heading_texts(doc)
        # Must not appear regardless of numbering
        assert not any("Comparison with Baseline Run" in h for h in headings)

    def test_comparison_present_when_available(self, tmp_path):
        run_dir = _make_run_dir(tmp_path, comparison=COMPARISON_AVAILABLE)
        out = str(tmp_path / "report.docx")
        result = _run_script(run_dir, out)
        assert result.returncode == 0, result.stderr
        doc = Document(out)
        headings = _heading_texts(doc)
        assert any("Comparison with Baseline Run" in h for h in headings)


class TestMissingAuditDir:
    """Missing 09_audit/ directory yields non-zero exit with clear error."""

    def test_missing_audit_exits_nonzero(self, tmp_path):
        run_dir = tmp_path / "bare_run"
        run_dir.mkdir()
        (run_dir / "manifest.json").write_text(json.dumps(MINIMAL_MANIFEST), encoding="utf-8")
        out = str(tmp_path / "report.docx")
        result = _run_script(str(run_dir), out)
        assert result.returncode != 0
        assert "09_audit" in result.stderr

    def test_missing_provenance_exits_nonzero(self, tmp_path):
        run_dir = tmp_path / "no_prov_run"
        run_dir.mkdir()
        audit_dir = run_dir / "09_audit"
        audit_dir.mkdir()
        (run_dir / "manifest.json").write_text(json.dumps(MINIMAL_MANIFEST), encoding="utf-8")
        out = str(tmp_path / "report.docx")
        result = _run_script(str(run_dir), out)
        assert result.returncode != 0
        assert "experiment_provenance.json" in result.stderr


class TestCustomTemplate:
    """Custom template with a subset of sections is honored."""

    def test_custom_two_section_template(self, tmp_path):
        custom_template = {
            "sections": [
                {"id": "cover", "title": "Custom Cover", "subtitle": "Test subtitle"},
                {"id": "run_summary", "title": "Summary Only"},
            ]
        }
        tmpl_path = str(tmp_path / "custom.yaml")
        with open(tmpl_path, "w", encoding="utf-8") as f:
            yaml.dump(custom_template, f)

        run_dir = _make_run_dir(tmp_path)
        out = str(tmp_path / "custom_report.docx")
        result = _run_script(run_dir, out, extra_args=["--template", tmpl_path])
        assert result.returncode == 0, result.stderr

        doc = Document(out)
        headings = _heading_texts(doc)
        assert "Custom Cover" in headings
        # run_summary is section 1 in this custom template
        assert "1 Summary Only" in headings
        # Sections not in the template must be absent
        assert not any("Artifact Provenance" in h for h in headings)
        assert not any("QA Gates" in h for h in headings)

    def test_json_template_also_works(self, tmp_path):
        custom_template = {
            "sections": [
                {"id": "cover", "title": "JSON Cover", "subtitle": ""},
                {"id": "appendix", "title": "Env Info"},
            ]
        }
        tmpl_path = str(tmp_path / "custom.json")
        with open(tmpl_path, "w", encoding="utf-8") as f:
            json.dump(custom_template, f)

        run_dir = _make_run_dir(tmp_path)
        out = str(tmp_path / "json_report.docx")
        result = _run_script(run_dir, out, extra_args=["--template", tmpl_path])
        assert result.returncode == 0, result.stderr
        doc = Document(out)
        headings = _heading_texts(doc)
        assert "JSON Cover" in headings
        assert "1 Env Info" in headings


class TestDeterminism:
    """Same inputs → identical paragraph+table text on two generations."""

    def test_content_determinism(self, tmp_path):
        run_dir = _make_run_dir(tmp_path, with_figures=True)
        out1 = str(tmp_path / "report1.docx")
        out2 = str(tmp_path / "report2.docx")

        r1 = _run_script(run_dir, out1)
        r2 = _run_script(run_dir, out2)
        assert r1.returncode == 0
        assert r2.returncode == 0

        doc1 = Document(out1)
        doc2 = Document(out2)

        texts1 = _all_text(doc1)
        texts2 = _all_text(doc2)
        assert texts1 == texts2, "Document text content differs between two generations"


class TestNoDatetimeNow:
    """The script must not call datetime.now() or time.time() — determinism lint."""

    def test_no_datetime_now_in_script(self):
        with open(SCRIPT, encoding="utf-8") as f:
            source = f.read()
        assert "datetime.now" not in source, "datetime.now() found in generate_report.py — forbidden"
        assert "time.time" not in source, "time.time() found in generate_report.py — forbidden"
        assert "datetime.utcnow" not in source, "datetime.utcnow() found in generate_report.py — forbidden"


# ---------------------------------------------------------------------------
# New engineering-format tests
# ---------------------------------------------------------------------------

class TestStyleColors:
    """Normal, Heading 1, Heading 2, Title styles must have pure black font color (000000)."""

    def test_normal_style_black(self, tmp_path):
        run_dir = _make_run_dir(tmp_path)
        out = str(tmp_path / "report.docx")
        result = _run_script(run_dir, out)
        assert result.returncode == 0, result.stderr
        doc = Document(out)
        color = doc.styles["Normal"].font.color.rgb
        assert str(color).upper() == "000000", f"Normal style color is {color}, expected 000000"

    def test_heading1_style_black(self, tmp_path):
        run_dir = _make_run_dir(tmp_path)
        out = str(tmp_path / "report.docx")
        result = _run_script(run_dir, out)
        assert result.returncode == 0, result.stderr
        doc = Document(out)
        color = doc.styles["Heading 1"].font.color.rgb
        assert str(color).upper() == "000000", f"Heading 1 style color is {color}, expected 000000"

    def test_heading2_style_black(self, tmp_path):
        run_dir = _make_run_dir(tmp_path)
        out = str(tmp_path / "report.docx")
        result = _run_script(run_dir, out)
        assert result.returncode == 0, result.stderr
        doc = Document(out)
        color = doc.styles["Heading 2"].font.color.rgb
        assert str(color).upper() == "000000", f"Heading 2 style color is {color}, expected 000000"

    def test_title_style_black(self, tmp_path):
        run_dir = _make_run_dir(tmp_path)
        out = str(tmp_path / "report.docx")
        result = _run_script(run_dir, out)
        assert result.returncode == 0, result.stderr
        doc = Document(out)
        color = doc.styles["Title"].font.color.rgb
        assert str(color).upper() == "000000", f"Title style color is {color}, expected 000000"


class TestTableCaptionsAndNarratives:
    """Every table is immediately preceded by a 'Table N — ' caption and followed by a narrative."""

    def _get_body_sequence(self, doc: Document):
        """Return list of ('para'|'table', text_or_None) tuples in body order."""
        items = []
        for kind, obj in _body_elements(doc):
            if kind == "para":
                items.append(("para", obj.text))
            else:
                items.append(("table", None))
        return items

    def test_every_table_preceded_by_caption(self, tmp_path):
        run_dir = _make_run_dir(tmp_path)
        out = str(tmp_path / "report.docx")
        result = _run_script(run_dir, out)
        assert result.returncode == 0, result.stderr
        doc = Document(out)
        seq = self._get_body_sequence(doc)

        caption_re = re.compile(r"^Table \d+ — ")
        for i, (kind, text) in enumerate(seq):
            if kind == "table":
                # Find the most recent non-empty paragraph before this table
                prev_para = None
                for j in range(i - 1, -1, -1):
                    if seq[j][0] == "para" and seq[j][1].strip():
                        prev_para = seq[j][1].strip()
                        break
                assert prev_para is not None, f"Table at index {i} has no preceding paragraph"
                assert caption_re.match(prev_para), (
                    f"Table at index {i} preceded by {prev_para!r}, "
                    f"expected 'Table N — ...' caption"
                )

    def test_every_table_followed_by_narrative(self, tmp_path):
        run_dir = _make_run_dir(tmp_path)
        out = str(tmp_path / "report.docx")
        result = _run_script(run_dir, out)
        assert result.returncode == 0, result.stderr
        doc = Document(out)
        seq = self._get_body_sequence(doc)

        for i, (kind, text) in enumerate(seq):
            if kind == "table":
                # Find the first non-empty paragraph after this table
                next_para = None
                for j in range(i + 1, len(seq)):
                    if seq[j][0] == "para" and seq[j][1].strip():
                        next_para = seq[j][1].strip()
                        break
                assert next_para is not None and len(next_para) > 0, (
                    f"Table at index {i} has no following narrative paragraph"
                )

    def test_table_caption_numbering_sequential(self, tmp_path):
        """Table captions must use sequential numbers starting at 1 with no gaps."""
        run_dir = _make_run_dir(tmp_path)
        out = str(tmp_path / "report.docx")
        result = _run_script(run_dir, out)
        assert result.returncode == 0, result.stderr
        doc = Document(out)

        caption_re = re.compile(r"^Table (\d+) — ")
        numbers = []
        for para in doc.paragraphs:
            m = caption_re.match(para.text.strip())
            if m:
                numbers.append(int(m.group(1)))

        assert len(numbers) >= 1, "No Table captions found"
        assert numbers == list(range(1, len(numbers) + 1)), (
            f"Table caption numbers are not sequential from 1: {numbers}"
        )


class TestFigureNumbering:
    """Figure captions use sequential 'Figure N — ' format (separate counter from tables)."""

    def test_figure_captions_numbered(self, tmp_path):
        run_dir = _make_run_dir(tmp_path, with_figures=True)
        out = str(tmp_path / "report.docx")
        result = _run_script(run_dir, out)
        assert result.returncode == 0, result.stderr
        doc = Document(out)

        fig_re = re.compile(r"^Figure (\d+) — ")
        numbers = []
        for para in doc.paragraphs:
            m = fig_re.match(para.text.strip())
            if m:
                numbers.append(int(m.group(1)))

        assert len(numbers) >= 2, f"Expected at least 2 figure captions, found: {numbers}"
        assert numbers == list(range(1, len(numbers) + 1)), (
            f"Figure caption numbers are not sequential from 1: {numbers}"
        )

    def test_figure_and_table_counters_independent(self, tmp_path):
        """Table N and Figure N counters must be separate — both start at 1."""
        run_dir = _make_run_dir(tmp_path, with_figures=True)
        out = str(tmp_path / "report.docx")
        result = _run_script(run_dir, out)
        assert result.returncode == 0, result.stderr
        doc = Document(out)

        tbl_re = re.compile(r"^Table (\d+) — ")
        fig_re = re.compile(r"^Figure (\d+) — ")
        tbl_nums = [int(m.group(1)) for p in doc.paragraphs if (m := tbl_re.match(p.text.strip()))]
        fig_nums = [int(m.group(1)) for p in doc.paragraphs if (m := fig_re.match(p.text.strip()))]

        assert tbl_nums[0] == 1, f"First table caption number should be 1, got {tbl_nums[0]}"
        assert fig_nums[0] == 1, f"First figure caption number should be 1, got {fig_nums[0]}"


class TestFooterPageField:
    """Footer must contain a Word PAGE field (fldChar + instrText ' PAGE ')."""

    def _footer_xml(self, doc: Document) -> str:
        section = doc.sections[0]
        footer = section.footer
        return footer._element.xml

    def test_footer_contains_page_field(self, tmp_path):
        run_dir = _make_run_dir(tmp_path)
        out = str(tmp_path / "report.docx")
        result = _run_script(run_dir, out)
        assert result.returncode == 0, result.stderr
        doc = Document(out)
        xml = self._footer_xml(doc)
        assert "PAGE" in xml, f"Footer XML does not contain PAGE field instruction. XML: {xml[:400]}"

    def test_footer_has_fldchar_begin(self, tmp_path):
        run_dir = _make_run_dir(tmp_path)
        out = str(tmp_path / "report.docx")
        result = _run_script(run_dir, out)
        assert result.returncode == 0, result.stderr
        doc = Document(out)
        xml = self._footer_xml(doc)
        assert "fldChar" in xml, f"Footer XML does not contain fldChar element. XML: {xml[:400]}"

    def test_footer_field_is_right_aligned(self, tmp_path):
        run_dir = _make_run_dir(tmp_path)
        out = str(tmp_path / "report.docx")
        result = _run_script(run_dir, out)
        assert result.returncode == 0, result.stderr
        doc = Document(out)
        section = doc.sections[0]
        footer = section.footer
        # The footer paragraph with the field should be right-aligned
        for para in footer.paragraphs:
            if para.text.strip() == "" and para._p.xml and "PAGE" in para._p.xml:
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                assert para.alignment == WD_ALIGN_PARAGRAPH.RIGHT or para.alignment is None, (
                    f"Footer paragraph alignment is {para.alignment}"
                )
                break


class TestNumberedSections:
    """Body sections are numbered sequentially; cover is not numbered."""

    def test_cover_heading_unnumbered(self, tmp_path):
        run_dir = _make_run_dir(tmp_path)
        out = str(tmp_path / "report.docx")
        result = _run_script(run_dir, out)
        assert result.returncode == 0, result.stderr
        doc = Document(out)
        headings = _heading_texts(doc)
        # Cover heading must not start with a digit
        assert "Run Audit Deliverable" in headings
        h1_headings = [h for h in headings if doc.paragraphs]
        cover_heading = headings[0]
        assert not cover_heading[0].isdigit(), (
            f"Cover heading should not be numbered, got: {cover_heading!r}"
        )

    def test_body_sections_numbered_sequentially(self, tmp_path):
        run_dir = _make_run_dir(tmp_path, comparison=COMPARISON_AVAILABLE)
        out = str(tmp_path / "report.docx")
        result = _run_script(run_dir, out)
        assert result.returncode == 0, result.stderr
        doc = Document(out)
        headings = _heading_texts(doc)

        # All headings starting with digits; extract section numbers
        numbered = [h for h in headings if h and h[0].isdigit()]
        nums = [int(h.split()[0]) for h in numbered]
        assert nums == list(range(1, len(nums) + 1)), (
            f"Section numbers are not sequential from 1: {nums} (headings: {numbered})"
        )

    def test_custom_template_renumbers_correctly(self, tmp_path):
        """A custom 3-section template must number body sections 1, 2 (cover exempt)."""
        custom_template = {
            "sections": [
                {"id": "cover", "title": "My Cover", "subtitle": ""},
                {"id": "run_summary", "title": "Summary"},
                {"id": "appendix", "title": "Env"},
            ]
        }
        tmpl_path = str(tmp_path / "custom.yaml")
        with open(tmpl_path, "w", encoding="utf-8") as f:
            yaml.dump(custom_template, f)

        run_dir = _make_run_dir(tmp_path)
        out = str(tmp_path / "custom_report.docx")
        result = _run_script(run_dir, out, extra_args=["--template", tmpl_path])
        assert result.returncode == 0, result.stderr
        doc = Document(out)
        headings = _heading_texts(doc)
        assert "1 Summary" in headings
        assert "2 Env" in headings
