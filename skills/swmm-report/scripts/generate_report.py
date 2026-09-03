#!/usr/bin/env python3
"""Generate a Word (.docx) deliverable from an audited SWMM run directory.

Usage
-----
    python generate_report.py --run-dir <path> [--out <path.docx>] [--template <yaml>]

Portability rule: this script imports ONLY stdlib + python-docx + PyYAML.
It must NEVER import agentic_swmm (skill-script portability requirement).
"""

from __future__ import annotations

import argparse
import glob
import json
import os
import re
import sys

# ---------------------------------------------------------------------------
# Guard optional dependency up-front so the error surfaces immediately.
# ---------------------------------------------------------------------------
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("python-docx is required: pip install 'aiswmm[report]'", file=sys.stderr)
    sys.exit(1)

try:
    import yaml
except ImportError:
    print("PyYAML is required: pip install PyYAML", file=sys.stderr)
    sys.exit(1)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _load_json(path: str) -> dict:
    """Load JSON; return empty dict if file absent."""
    if os.path.exists(path):
        with open(path, encoding="utf-8") as f:
            return json.load(f)
    return {}


def _load_template(template_path: str) -> dict:
    """Load a YAML (or JSON) template file by extension."""
    ext = os.path.splitext(template_path)[1].lower()
    with open(template_path, encoding="utf-8") as f:
        if ext in (".yaml", ".yml"):
            return yaml.safe_load(f)
        elif ext == ".json":
            return json.load(f)
        else:
            raise ValueError(f"Unsupported template extension: {ext!r}. Use .yaml or .json.")


def _default_template_path() -> str:
    return os.path.join(os.path.dirname(__file__), "..", "templates", "default.yaml")


def _continuity_text(value) -> str:
    """Render the continuity metric whatever shape the audit wrote.

    The audit's provenance carries a structured metric
    (``{"values": {"flow_routing": 0.402, "runoff_quantity": -0.09}, ...}``);
    the report printed that dict verbatim in a client deliverable (live
    finding F-17, 2026-09-02). A bare number and a missing value still
    render as before.
    """
    if isinstance(value, dict):
        values = value.get("values")
        if isinstance(values, dict) and values:
            parts = []
            for key, label in (("flow_routing", "routing"), ("runoff_quantity", "runoff")):
                if key in values and values[key] is not None:
                    parts.append(f"{label} {values[key]}")
            for key, number in values.items():
                if key not in ("flow_routing", "runoff_quantity") and number is not None:
                    parts.append(f"{key} {number}")
            if parts:
                return " / ".join(parts)
        if value.get("value") is not None:
            return _na(value.get("value"))
        return "n/a"
    return _na(value)


def _na(value) -> str:
    """Return 'n/a' for None/missing values."""
    if value is None:
        return "n/a"
    return str(value)


def _set_style_black(doc: Document) -> None:
    """Override font color to pure black for all styles used by this generator."""
    black = RGBColor(0, 0, 0)
    for style_name in ("Normal", "Heading 1", "Heading 2", "Title"):
        try:
            style = doc.styles[style_name]
            style.font.color.rgb = black
        except KeyError:
            pass  # Style absent in this template — skip silently


def _add_table_caption(doc: Document, caption_text: str, table_counter: list) -> None:
    """Insert an engineering-convention table caption ABOVE the next table.

    Format: bold ``Table N`` prefix followed by em-dash and caption text, all black.
    ``table_counter`` is a one-element list used as a mutable integer.
    """
    table_counter[0] += 1
    n = table_counter[0]
    cap_para = doc.add_paragraph()
    # Bold "Table N:" prefix (house style: no em dashes in deliverables)
    bold_run = cap_para.add_run(f"Table {n}: ")
    bold_run.bold = True
    bold_run.font.color.rgb = RGBColor(0, 0, 0)
    # Regular caption text
    text_run = cap_para.add_run(caption_text)
    text_run.bold = False
    text_run.font.color.rgb = RGBColor(0, 0, 0)


def _add_narrative(doc: Document, narrative_text: str) -> None:
    """Insert an explanatory narrative paragraph immediately after a table."""
    p = doc.add_paragraph(narrative_text)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)


def _add_figure_caption(doc: Document, caption_text: str, figure_counter: list) -> None:
    """Insert an engineering-convention figure caption BELOW the figure.

    Format: ``Figure N: <text>``, centred, all black.
    ``figure_counter`` is a one-element list used as a mutable integer.
    """
    figure_counter[0] += 1
    n = figure_counter[0]
    cap_para = doc.add_paragraph(f"Figure {n}: {caption_text}")
    cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap_para.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)


def _add_page_number_footer(doc: Document) -> None:
    """Add a right-aligned PAGE field to the default section footer.

    python-docx has no high-level PAGE-field API; we insert the standard OOXML
    field via fldChar + instrText runs so Word/LibreOffice renders a live number.
    """
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False

    # Clear any existing paragraphs and work with the first one
    for para in footer.paragraphs:
        # Remove all runs from existing paragraph
        p_elem = para._p
        for child in list(p_elem):
            if child.tag != qn("w:pPr"):
                p_elem.remove(child)

    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Build the PAGE field via raw OOXML: fldChar(begin) + instrText + fldChar(end)
    run_begin = footer_para.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run_begin._r.append(fld_begin)

    run_instr = footer_para.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    run_instr._r.append(instr)

    run_end = footer_para.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run_end._r.append(fld_end)


def _numbered_heading(section_number: int, title: str) -> str:
    """Return ``N Title`` for body sections (section_number >= 1)."""
    return f"{section_number} {title}"


# ---------------------------------------------------------------------------
# Section renderers
# Each renderer has signature: (doc, section_cfg, artifacts, ctx) -> None
# `artifacts` is a dict carrying the loaded JSON data and discovered PNG paths.
# `ctx` is a mutable rendering context dict:
#   ctx["table_counter"]: [int]  — auto-incrementing table number
#   ctx["figure_counter"]: [int] — auto-incrementing figure number
#   ctx["section_number"]: int   — increments for each body heading (cover exempt)
# ---------------------------------------------------------------------------

def _render_cover(doc: Document, cfg: dict, artifacts: dict, ctx: dict) -> None:
    prov = artifacts.get("provenance", {})
    manifest = artifacts.get("manifest", {})

    title = cfg.get("title", "Run Audit Deliverable")
    subtitle = cfg.get("subtitle", "")

    # Cover title is unnumbered (level 1 heading, centred)
    heading = doc.add_heading(title, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if subtitle:
        p = doc.add_paragraph(subtitle)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run_id = prov.get("run_id") or manifest.get("run_dir", "")
    generated_at = prov.get("generated_at_utc", "")
    swmm_version = (prov.get("tools") or {}).get("swmm5_version", "")

    caption = cfg.get("caption", "Run identification and generation metadata.")
    narrative = cfg.get(
        "narrative",
        "Run ID and timestamps are sourced directly from experiment_provenance.json; "
        "the SWMM version reflects the binary used during the simulation run.",
    )

    rows_data = [
        ("Run ID", _na(run_id)),
        ("Generated at (UTC)", _na(generated_at)),
        ("SWMM version", _na(swmm_version)),
    ]

    _add_table_caption(doc, caption, ctx["table_counter"])
    info_table = doc.add_table(rows=3, cols=2)
    info_table.style = "Table Grid"
    for row, (label, value) in zip(info_table.rows, rows_data):
        row.cells[0].text = label
        row.cells[1].text = value

    _add_narrative(doc, narrative)
    doc.add_paragraph()


def _render_run_summary(doc: Document, cfg: dict, artifacts: dict, ctx: dict) -> None:
    prov = artifacts.get("provenance", {})
    manifest = artifacts.get("manifest", {})

    ctx["section_number"] += 1
    raw_title = cfg.get("title", "Run Summary")
    doc.add_heading(_numbered_heading(ctx["section_number"], raw_title), level=2)

    metrics = prov.get("metrics", {})
    peak_flow_obj = metrics.get("peak_flow") or {}
    peak_flow_val = peak_flow_obj.get("value") if isinstance(peak_flow_obj, dict) else None
    peak_flow_unit = peak_flow_obj.get("unit") if isinstance(peak_flow_obj, dict) else None
    time_of_peak = peak_flow_obj.get("time_hhmm") if isinstance(peak_flow_obj, dict) else None
    continuity_error = metrics.get("continuity_error")
    return_code = metrics.get("swmm_return_code")

    # Fall back to manifest qoi when provenance metrics absent
    if peak_flow_val is None:
        qoi = manifest.get("qoi", {})
        peak_flow_val = qoi.get("peak_flow_cms_at_O1")
        peak_flow_unit = "CMS"
        time_of_peak = qoi.get("time_of_peak_hhmm")

    columns = cfg.get("columns", ["Metric", "Value", "Unit"])
    rows_cfg = cfg.get("rows", [])

    caption = cfg.get("caption", "Key hydraulic performance metrics for this simulation run.")
    narrative = cfg.get(
        "narrative",
        "Values are read from the run manifest and experiment_provenance.json produced by "
        "the SWMM runner and audit tool; no metrics are recomputed at report generation time.",
    )

    value_map = {
        "peak_flow": (_na(peak_flow_val), peak_flow_unit or "flow units not recorded"),
        "time_of_peak": (_na(time_of_peak), "hh:mm"),
        "continuity_error": (_continuity_text(continuity_error), "%"),
        "return_code": (_na(return_code), ""),
    }

    _add_table_caption(doc, caption, ctx["table_counter"])
    table = doc.add_table(rows=1 + len(rows_cfg), cols=len(columns))
    table.style = "Table Grid"
    header_row = table.rows[0]
    for i, col_name in enumerate(columns):
        header_row.cells[i].text = col_name

    for i, row_cfg in enumerate(rows_cfg):
        key = row_cfg.get("key", "")
        label = row_cfg.get("label", key)
        unit = row_cfg.get("unit", "")
        val, default_unit = value_map.get(key, ("n/a", unit))
        data_row = table.rows[i + 1]
        data_row.cells[0].text = label
        data_row.cells[1].text = val
        if len(columns) > 2:
            data_row.cells[2].text = unit or default_unit

    _add_narrative(doc, narrative)
    doc.add_paragraph()


def _render_model_description(doc: Document, cfg: dict, artifacts: dict, ctx: dict) -> None:
    manifest = artifacts.get("manifest", {})
    inp_facts = artifacts.get("inp_facts", {}) or {}

    ctx["section_number"] += 1
    raw_title = cfg.get("title", "Model Description")
    doc.add_heading(_numbered_heading(ctx["section_number"], raw_title), level=2)

    # Source priority: facts parsed from the run's own INP (present for
    # every run), then legacy manifest keys (one-off benchmark runs
    # that hand-wrote them keep rendering), then n/a. The historical
    # manifest-only read rendered n/a for every production run because
    # no production writer emits these keys.
    landuse_params = manifest.get("landuse_params", {}) or {}
    green_ampt = manifest.get("green_ampt_params", {}) or {}

    def _fact(key: str, legacy_value: object) -> str:
        value = inp_facts.get(key)
        if value is None:
            value = legacy_value
        return _na(value)

    value_map = {
        "basin_area_ha": _fact("basin_area_ha", manifest.get("basin_area_ha")),
        "sim_start": _fact("sim_start", manifest.get("sim_start")),
        "sim_end": _fact("sim_end", manifest.get("sim_end")),
        "imperv_pct": _fact("imperv_pct", landuse_params.get("imperv_pct")),
        "ksat_mm_per_hr": _fact("ksat_mm_per_hr", green_ampt.get("ksat_mm_per_hr")),
        "suction_mm": _fact("suction_mm", green_ampt.get("suction_mm")),
    }

    columns = cfg.get("columns", ["Parameter", "Value"])
    rows_cfg = cfg.get("rows", [])

    caption = cfg.get("caption", "Catchment geometry, simulation window, and infiltration parameters.")
    narrative = cfg.get(
        "narrative",
        "Geometry and simulation-window values are parsed from the run's own SWMM INP "
        "(total subcatchment area, area-weighted imperviousness, [OPTIONS] dates); "
        "infiltration values are subcatchment averages and are reported only when the "
        "model uses Green-Ampt. Values shown as n/a could not be derived from the run "
        "directory.",
    )

    _add_table_caption(doc, caption, ctx["table_counter"])
    table = doc.add_table(rows=1 + len(rows_cfg), cols=len(columns))
    table.style = "Table Grid"
    header_row = table.rows[0]
    for i, col_name in enumerate(columns):
        header_row.cells[i].text = col_name

    for i, row_cfg in enumerate(rows_cfg):
        key = row_cfg.get("key", "")
        label = row_cfg.get("label", key)
        unit = row_cfg.get("unit", "")
        val = value_map.get(key, "n/a")
        data_row = table.rows[i + 1]
        data_row.cells[0].text = f"{label} ({unit})" if unit else label
        data_row.cells[1].text = val

    _add_narrative(doc, narrative)
    doc.add_paragraph()


def _render_qa_gates(doc: Document, cfg: dict, artifacts: dict, ctx: dict) -> None:
    prov = artifacts.get("provenance", {})

    ctx["section_number"] += 1
    raw_title = cfg.get("title", "QA Gates")
    doc.add_heading(_numbered_heading(ctx["section_number"], raw_title), level=2)

    qa = prov.get("qa", {})
    checks = qa.get("checks", [])
    pass_label = cfg.get("pass_label", "PASS")
    fail_label = cfg.get("fail_label", "FAIL")

    columns = cfg.get("columns", ["Gate ID", "Status", "Detail"])

    caption = cfg.get("caption", "Automated QA gate results from the audit pipeline.")
    narrative = cfg.get(
        "narrative",
        "Each gate is evaluated by swmm-experiment-audit against thresholds defined in "
        "the project configuration; PASS/FAIL status is read from experiment_provenance.json "
        "and is not recomputed here.",
    )

    if checks:
        _add_table_caption(doc, caption, ctx["table_counter"])
        table = doc.add_table(rows=1 + len(checks), cols=len(columns))
        table.style = "Table Grid"
        header_row = table.rows[0]
        for i, col_name in enumerate(columns):
            header_row.cells[i].text = col_name

        for i, check in enumerate(checks):
            data_row = table.rows[i + 1]
            data_row.cells[0].text = _na(check.get("id"))
            status = pass_label if check.get("ok") else fail_label
            data_row.cells[1].text = status
            if len(columns) > 2:
                data_row.cells[2].text = _na(check.get("detail"))

        _add_narrative(doc, narrative)
    else:
        p = doc.add_paragraph("No QA checks recorded.")
        p.italic = True

    doc.add_paragraph()


def _render_design_review(doc: Document, cfg: dict, artifacts: dict, ctx: dict) -> None:
    """The rulebook verdict, per rule, with the worst element.

    Live finding F-18 (2026-09-02): the shell's final answer said "Design
    review: FAIL (1 pass, 2 fail, 4 warn, 4 needs-data)" while the Word
    deliverable carried no review at all. The client reads the deliverable.
    """
    review = artifacts.get("design_review", {}) or {}

    ctx["section_number"] += 1
    raw_title = cfg.get("title", "Design Review")
    doc.add_heading(_numbered_heading(ctx["section_number"], raw_title), level=2)

    results = review.get("results") if isinstance(review, dict) else None
    if not review or not isinstance(results, list) or not results:
        p = doc.add_paragraph(cfg.get("missing_text", "No design review recorded for this run."))
        p.italic = True
        doc.add_paragraph()
        return

    summary = review.get("summary", {}) if isinstance(review.get("summary"), dict) else {}
    overall = str(review.get("overall_status", "n/a")).upper()
    verdict = cfg.get(
        "verdict_text",
        "Overall result: {overall} ({passed} pass, {failed} fail, {warned} warn, "
        "{needs_data} needs-data) against rulebook {rulebook} v{version}.",
    ).format(
        overall=overall,
        passed=summary.get("pass", 0),
        failed=summary.get("fail", 0),
        warned=summary.get("warn", 0),
        needs_data=summary.get("needs_data", 0),
        rulebook=review.get("rulebook_id", "n/a"),
        version=review.get("rulebook_version", "n/a"),
    )
    doc.add_paragraph(verdict)

    order = {"fail": 0, "warn": 1, "pass": 2, "needs_data": 3}
    rows = sorted(results, key=lambda r: order.get(str(r.get("status", "")).lower(), 9))
    columns = cfg.get("columns", ["Rule", "Status", "Worst element", "Value", "Threshold"])
    caption = cfg.get("caption", "Design-review rules and the element that decided each verdict.")
    _add_table_caption(doc, caption, ctx["table_counter"])
    table = doc.add_table(rows=1 + len(rows), cols=len(columns))
    table.style = "Table Grid"
    for i, col_name in enumerate(columns):
        table.rows[0].cells[i].text = col_name
    for i, rule in enumerate(rows):
        worst = rule.get("worst_element") if isinstance(rule.get("worst_element"), dict) else {}
        cells = table.rows[i + 1].cells
        cells[0].text = _na(rule.get("title") or rule.get("rule_id"))
        cells[1].text = str(rule.get("status", "n/a")).upper()
        if len(columns) > 2:
            cells[2].text = _na(worst.get("id")) if worst else (
                _na(rule.get("needs_data_reason")) if rule.get("needs_data_reason") else "n/a"
            )
        if len(columns) > 3:
            cells[3].text = _na(worst.get("value")) if worst else "n/a"
        if len(columns) > 4:
            cells[4].text = _na(worst.get("threshold")) if worst else "n/a"
    narrative = cfg.get(
        "narrative",
        "Statuses are read from 11_review/design_review.json as written by the design-review "
        "tool; nothing is re-evaluated at report generation time.",
    )
    _add_narrative(doc, narrative)
    disclaimer = review.get("disclaimer")
    if disclaimer:
        p = doc.add_paragraph(str(disclaimer).replace("\n", " "))
        p.italic = True
    doc.add_paragraph()


def _render_evidence_boundary(doc: Document, cfg: dict, artifacts: dict, ctx: dict) -> None:
    """What these numbers are, and what they are not.

    Live finding F-18 (2026-09-02): "uncalibrated" appeared zero times in a
    client deliverable built from an uncalibrated first-pass model, while
    the shell's answer led with that boundary. The deliverable now states
    the calibration status and the review verdict in words.
    """
    prov = artifacts.get("provenance", {}) or {}
    review = artifacts.get("design_review", {}) or {}

    ctx["section_number"] += 1
    raw_title = cfg.get("title", "Evidence Boundary")
    doc.add_heading(_numbered_heading(ctx["section_number"], raw_title), level=2)

    calibration = prov.get("calibration")
    status = None
    if isinstance(calibration, dict):
        status = calibration.get("status") or calibration.get("calibration_status")
    elif isinstance(calibration, str):
        status = calibration
    if status and str(status).lower() not in ("none", "uncalibrated", "not_calibrated", "n/a"):
        doc.add_paragraph(
            cfg.get(
                "calibrated_text",
                "Calibration status recorded for this run: {status}. Results are only as "
                "trustworthy as the observed data and objective behind that status.",
            ).format(status=status)
        )
    else:
        doc.add_paragraph(
            cfg.get(
                "uncalibrated_text",
                "This is an uncalibrated first-pass model. No observed flow data were used, "
                "so every value in this report is a simulated result, not a validated "
                "prediction. Calibrate against observed flow before design decisions.",
            )
        )

    if isinstance(review, dict) and review.get("overall_status"):
        summary = review.get("summary", {}) if isinstance(review.get("summary"), dict) else {}
        doc.add_paragraph(
            cfg.get(
                "review_text",
                "The design review against rulebook {rulebook} returned {overall} "
                "({failed} fail, {warned} warn). The bundled rulebook is a template: it "
                "does not certify compliance with any adopted standard.",
            ).format(
                rulebook=review.get("rulebook_id", "n/a"),
                overall=str(review.get("overall_status", "n/a")).upper(),
                failed=summary.get("fail", 0),
                warned=summary.get("warn", 0),
            )
        )
    else:
        doc.add_paragraph(
            cfg.get("no_review_text", "No design review was run against this model.")
        )
    doc.add_paragraph()


def _render_figures(doc: Document, cfg: dict, artifacts: dict, ctx: dict) -> None:
    run_dir = artifacts.get("run_dir", "")

    ctx["section_number"] += 1
    raw_title = cfg.get("title", "Figures")
    doc.add_heading(_numbered_heading(ctx["section_number"], raw_title), level=2)

    # Canonical stage first (08_plot per the run layout), then both
    # legacy generations: 07_plots (generation-B CLI runs, and the name
    # a planner freehand-writes most often) and 07_plot (the original
    # template default). Scanning one name while figures land in
    # another silently produced figureless reports (found 2026-08-09).
    plot_dirs = cfg.get("plot_dirs", ["00_raw", "08_plot", "08_plots", "07_plots", "07_plot"])
    include_root_png = cfg.get("include_root_png", "network_layout.png")
    no_figures_note = cfg.get("no_figures_note", "No figures available for this run.")

    png_paths = []
    for plot_dir in plot_dirs:
        pattern = os.path.join(run_dir, plot_dir, "*.png")
        png_paths.extend(sorted(glob.glob(pattern)))

    # Include optional root-level PNG (e.g. network_layout.png)
    if include_root_png:
        root_png = os.path.join(run_dir, include_root_png)
        if os.path.exists(root_png) and root_png not in png_paths:
            png_paths.append(root_png)

    if not png_paths:
        p = doc.add_paragraph(no_figures_note)
        p.italic = True
    else:
        for idx, png_path in enumerate(png_paths, start=1):
            stem = os.path.splitext(os.path.basename(png_path))[0]
            try:
                doc.add_picture(png_path, width=Inches(5.5))
            except Exception:
                doc.add_paragraph(f"[image {idx}: {stem}, could not be embedded]")
            # Caption BELOW figure (engineering convention), numbered
            _add_figure_caption(doc, stem, ctx["figure_counter"])
            doc.add_paragraph()

    doc.add_paragraph()


def _render_diagnostics(doc: Document, cfg: dict, artifacts: dict, ctx: dict) -> None:
    diag_data = artifacts.get("diagnostics", {})

    ctx["section_number"] += 1
    raw_title = cfg.get("title", "Model Diagnostics")
    doc.add_heading(_numbered_heading(ctx["section_number"], raw_title), level=2)

    diagnostics = diag_data.get("diagnostics", [])
    no_diag_note = cfg.get("no_diagnostics_note", "No diagnostics reported.")
    columns = cfg.get("columns", ["Code", "Severity", "Message"])

    caption = cfg.get("caption", "SWMM diagnostic messages emitted during the simulation run.")
    narrative = cfg.get(
        "narrative",
        "Diagnostic entries are parsed from model_diagnostics.json produced by "
        "swmm-experiment-audit; codes and severities follow the SWMM 5 output format.",
    )

    if not diagnostics:
        # Also report overall status
        status = diag_data.get("status", "unknown")
        error_count = diag_data.get("error_count", 0)
        warning_count = diag_data.get("warning_count", 0)
        p = doc.add_paragraph(
            f"{no_diag_note} Overall status: {status.upper()} "
            f"(errors: {error_count}, warnings: {warning_count})."
        )
        p.italic = True
    else:
        _add_table_caption(doc, caption, ctx["table_counter"])
        table = doc.add_table(rows=1 + len(diagnostics), cols=len(columns))
        table.style = "Table Grid"
        header_row = table.rows[0]
        for i, col_name in enumerate(columns):
            header_row.cells[i].text = col_name
        for i, diag in enumerate(diagnostics):
            data_row = table.rows[i + 1]
            data_row.cells[0].text = _na(diag.get("code"))
            if len(columns) > 1:
                data_row.cells[1].text = _na(diag.get("severity"))
            if len(columns) > 2:
                data_row.cells[2].text = _na(diag.get("message") or diag.get("msg"))

        _add_narrative(doc, narrative)

    doc.add_paragraph()


def _render_hydraulic_results(doc: Document, cfg: dict, artifacts: dict, ctx: dict) -> None:
    """Node, outfall and conduit results from the run's own .rpt.

    Without this the deliverable carried provenance and QA gates and no
    hydraulics, so a reader asked where the node flows were. SWMM had computed
    them; nothing downstream looked.
    """
    data = artifacts.get("hydraulics") or {}

    ctx["section_number"] += 1
    raw_title = cfg.get("title", "Hydraulic Results")
    doc.add_heading(_numbered_heading(ctx["section_number"], raw_title), level=2)

    if not data:
        p = doc.add_paragraph(
            cfg.get(
                "no_results_note",
                "No hydraulic summary was extracted for this run, so no node, "
                "outfall or conduit results are reported here.",
            )
        )
        p.italic = True
        doc.add_paragraph()
        return

    units = data.get("flow_units") or "flow units per model"
    counts = data.get("counts") or {}
    top_n = data.get("top_n")

    def _shown_of(kind: str) -> str:
        total = counts.get(kind)
        rows = data.get(kind) or []
        if total and total > len(rows):
            return f" (top {len(rows)} of {total} by peak)"
        return ""

    nodes = data.get("nodes") or []
    if nodes:
        _add_table_caption(
            doc,
            cfg.get("nodes_caption", f"Node inflow summary, flow in {units}{_shown_of('nodes')}."),
            ctx["table_counter"],
        )
        columns = ["Node", "Type", f"Peak total inflow ({units})", "Time of peak", "Flow balance error (%)"]
        table = doc.add_table(rows=1 + len(nodes), cols=len(columns))
        table.style = "Table Grid"
        for i, name in enumerate(columns):
            table.rows[0].cells[i].text = name
        for i, row in enumerate(nodes):
            cells = table.rows[i + 1].cells
            cells[0].text = _na(row.get("node"))
            cells[1].text = _na(row.get("type"))
            cells[2].text = _na(row.get("max_total_inflow"))
            cells[3].text = _na(row.get("time_of_max"))
            cells[4].text = _na(row.get("flow_balance_error_pct"))
        doc.add_paragraph()

    outfalls = data.get("outfalls") or []
    if outfalls:
        _add_table_caption(
            doc,
            cfg.get("outfalls_caption", f"Outfall loading summary, flow in {units}{_shown_of('outfalls')}."),
            ctx["table_counter"],
        )
        columns = ["Outfall", "Flow frequency (%)", f"Average flow ({units})", f"Peak flow ({units})", "Total volume (10^6 L)"]
        table = doc.add_table(rows=1 + len(outfalls), cols=len(columns))
        table.style = "Table Grid"
        for i, name in enumerate(columns):
            table.rows[0].cells[i].text = name
        for i, row in enumerate(outfalls):
            cells = table.rows[i + 1].cells
            cells[0].text = _na(row.get("node"))
            cells[1].text = _na(row.get("flow_freq_pct"))
            cells[2].text = _na(row.get("avg_flow"))
            cells[3].text = _na(row.get("max_flow"))
            cells[4].text = _na(row.get("total_volume_10_6_ltr"))
        doc.add_paragraph()

    links = data.get("links") or []
    if links:
        _add_table_caption(
            doc,
            cfg.get("links_caption", f"Conduit peak flows, flow in {units}{_shown_of('links')}."),
            ctx["table_counter"],
        )
        columns = ["Link", "Type", f"Peak flow ({units})", "Time of peak", "Max/full flow", "Max/full depth"]
        table = doc.add_table(rows=1 + len(links), cols=len(columns))
        table.style = "Table Grid"
        for i, name in enumerate(columns):
            table.rows[0].cells[i].text = name
        for i, row in enumerate(links):
            cells = table.rows[i + 1].cells
            cells[0].text = _na(row.get("link"))
            cells[1].text = _na(row.get("type"))
            cells[2].text = _na(row.get("peak_flow"))
            cells[3].text = _na(f"{row.get('time_days')} {row.get('time_hhmm')}" if row.get("time_hhmm") else None)
            cells[4].text = _na(row.get("max_full_flow_ratio"))
            cells[5].text = _na(row.get("max_full_depth_ratio"))
        doc.add_paragraph()

    _add_narrative(
        doc,
        cfg.get(
            "narrative",
            "Values are read from the run's own SWMM report file. A completed "
            "simulation is not a calibrated or validated one: these are model "
            "outputs, not measurements.",
        ),
    )
    doc.add_paragraph()


def _render_comparison(doc: Document, cfg: dict, artifacts: dict, ctx: dict) -> None:
    """Conditionally rendered — skipped entirely when comparison_available is false."""
    comp = artifacts.get("comparison", {})
    if not comp.get("comparison_available", False):
        return  # Silent skip — no heading added

    ctx["section_number"] += 1
    raw_title = cfg.get("title", "Comparison with Baseline Run")
    doc.add_heading(_numbered_heading(ctx["section_number"], raw_title), level=2)

    columns = cfg.get("columns", ["Metric", "Current", "Baseline", "Delta"])
    comp_metrics = comp.get("metrics", {}) or {}

    caption = cfg.get("caption", "Side-by-side metric comparison between the current and baseline runs.")
    narrative = cfg.get(
        "narrative",
        "Current and baseline values are read from 09_audit/comparison.json generated by "
        "swmm-experiment-audit; delta is the arithmetic difference (current minus baseline) "
        "as stored in that file and is not recomputed here.",
    )

    rows = []
    for metric_key, metric_obj in comp_metrics.items():
        if isinstance(metric_obj, dict):
            current = metric_obj.get("current", "n/a")
            baseline = metric_obj.get("baseline", "n/a")
            delta = metric_obj.get("delta", "n/a")
            rows.append((metric_key, _na(current), _na(baseline), _na(delta)))
        else:
            rows.append((metric_key, _na(metric_obj), "n/a", "n/a"))

    if rows:
        _add_table_caption(doc, caption, ctx["table_counter"])
        table = doc.add_table(rows=1 + len(rows), cols=min(len(columns), 4))
        table.style = "Table Grid"
        header_row = table.rows[0]
        for i, col_name in enumerate(columns[:4]):
            header_row.cells[i].text = col_name
        for i, (mk, cur, bas, dlt) in enumerate(rows):
            data_row = table.rows[i + 1]
            data_row.cells[0].text = mk
            if len(columns) > 1:
                data_row.cells[1].text = cur
            if len(columns) > 2:
                data_row.cells[2].text = bas
            if len(columns) > 3:
                data_row.cells[3].text = dlt
        _add_narrative(doc, narrative)
    else:
        doc.add_paragraph("Comparison data is available but contains no metric rows.")

    doc.add_paragraph()


def _render_provenance(doc: Document, cfg: dict, artifacts: dict, ctx: dict) -> None:
    prov = artifacts.get("provenance", {})

    ctx["section_number"] += 1
    raw_title = cfg.get("title", "Artifact Provenance")
    doc.add_heading(_numbered_heading(ctx["section_number"], raw_title), level=2)

    art_dict = prov.get("artifacts", {})
    sha256_length = cfg.get("sha256_length", 16)
    columns = cfg.get("columns", ["Artifact ID", "Role", f"SHA-256 (first {sha256_length} chars)", "Produced by"])

    caption = cfg.get("caption", "Cryptographic hashes and roles of artifacts involved in this run.")
    narrative = cfg.get(
        "narrative",
        "SHA-256 digests are computed by swmm-experiment-audit at audit time and stored in "
        "experiment_provenance.json; only the first 16 hex characters are shown here for "
        "readability. The full digest is available in the provenance JSON.",
    )

    existing = [(k, v) for k, v in art_dict.items() if v.get("exists")]

    if not existing:
        doc.add_paragraph("No artifact hashes available.")
    else:
        _add_table_caption(doc, caption, ctx["table_counter"])
        table = doc.add_table(rows=1 + len(existing), cols=len(columns))
        table.style = "Table Grid"
        header_row = table.rows[0]
        for i, col_name in enumerate(columns):
            header_row.cells[i].text = col_name
        for i, (art_id, art_info) in enumerate(existing):
            sha = art_info.get("sha256") or ""
            sha_short = sha[:sha256_length] if sha else "n/a"
            data_row = table.rows[i + 1]
            data_row.cells[0].text = art_id
            if len(columns) > 1:
                data_row.cells[1].text = _na(art_info.get("role"))
            if len(columns) > 2:
                data_row.cells[2].text = sha_short
            if len(columns) > 3:
                data_row.cells[3].text = _na(art_info.get("produced_by"))

        _add_narrative(doc, narrative)

    doc.add_paragraph()


def _render_appendix(doc: Document, cfg: dict, artifacts: dict, ctx: dict) -> None:
    prov = artifacts.get("provenance", {})

    ctx["section_number"] += 1
    raw_title = cfg.get("title", "Appendix: Generation Environment")
    doc.add_heading(_numbered_heading(ctx["section_number"], raw_title), level=2)

    repo = prov.get("repo", {}) or {}
    tools = prov.get("tools", {}) or {}
    schema_version = prov.get("schema_version", "")
    generated_by = prov.get("generated_by", "")

    columns = cfg.get("columns", ["Item", "Value"])
    rows_data = [
        ("Git head", _na(repo.get("git_head"))),
        ("Git branch", _na(repo.get("git_branch"))),
        ("SWMM version", _na(tools.get("swmm5_version"))),
        ("Python version", _na(tools.get("python_version"))),
        ("Provenance schema version", _na(schema_version)),
        ("Generated by", _na(generated_by)),
    ]

    caption = cfg.get("caption", "Software environment at the time of audit generation.")
    narrative = cfg.get(
        "narrative",
        "All environment values are read from experiment_provenance.json; they capture the "
        "exact tool versions used to produce and audit the simulation results.",
    )

    _add_table_caption(doc, caption, ctx["table_counter"])
    table = doc.add_table(rows=1 + len(rows_data), cols=len(columns))
    table.style = "Table Grid"
    header_row = table.rows[0]
    for i, col_name in enumerate(columns):
        header_row.cells[i].text = col_name
    for i, (label, value) in enumerate(rows_data):
        data_row = table.rows[i + 1]
        data_row.cells[0].text = label
        if len(columns) > 1:
            data_row.cells[1].text = value

    _add_narrative(doc, narrative)
    doc.add_paragraph()


# ---------------------------------------------------------------------------
# Section registry — closed vocabulary
# ---------------------------------------------------------------------------

SECTION_RENDERERS = {
    "cover": _render_cover,
    "run_summary": _render_run_summary,
    "model_description": _render_model_description,
    "qa_gates": _render_qa_gates,
    "design_review": _render_design_review,
    "evidence_boundary": _render_evidence_boundary,
    "hydraulic_results": _render_hydraulic_results,
    "figures": _render_figures,
    "diagnostics": _render_diagnostics,
    "comparison": _render_comparison,
    "provenance": _render_provenance,
    "appendix": _render_appendix,
}


# ---------------------------------------------------------------------------
# Artifact loader
# ---------------------------------------------------------------------------

def _find_run_inp(run_dir: str) -> str | None:
    """Locate the run's INP: builder stage first, then inputs, then flat."""
    candidates = []
    for sub in ("05_builder", "04_builder", "00_inputs"):
        d = os.path.join(run_dir, sub)
        if os.path.isdir(d):
            candidates.extend(
                sorted(
                    os.path.join(d, name)
                    for name in os.listdir(d)
                    if name.lower().endswith(".inp")
                )
            )
    candidates.extend(
        sorted(
            os.path.join(run_dir, name)
            for name in os.listdir(run_dir)
            if name.lower().endswith(".inp")
        )
    )
    return candidates[0] if candidates else None


def _split_inp_sections(text: str) -> dict[str, list[str]]:
    """Return ``{SECTION_NAME: data_lines}`` (comments/blank stripped)."""
    sections: dict[str, list[str]] = {}
    current: list[str] | None = None
    for raw in text.splitlines():
        line = raw.strip()
        if not line or line.startswith(";"):
            continue
        m = re.match(r"^\[([A-Za-z_ ]+)\]$", line)
        if m:
            current = sections.setdefault(m.group(1).strip().upper(), [])
            continue
        if current is not None:
            current.append(line)
    return sections


def _inp_model_facts(run_dir: str) -> dict:
    """Derive model-description facts from the run's own INP.

    Every run carries its INP; the historical code instead read
    ``basin_area_ha`` / ``sim_start`` / ``landuse_params`` /
    ``green_ampt_params`` from the top manifest, a schema no production
    writer emits (only a one-off benchmark script did), so the default
    report's Model Description table rendered ``n/a`` throughout while
    its narrative claimed exact inputs (found 2026-08-08 drift sweep).

    Returned keys mirror the template's row keys. Values that cannot
    be derived (e.g. Green-Ampt columns under a Horton model) are
    simply absent — callers fall back to legacy manifest keys, then
    ``n/a``.
    """
    inp_path = _find_run_inp(run_dir)
    if inp_path is None:
        return {}
    try:
        text = open(inp_path, "r", encoding="utf-8", errors="replace").read()
    except OSError:
        return {}
    sections = _split_inp_sections(text)
    facts: dict[str, object] = {}

    options: dict[str, str] = {}
    for line in sections.get("OPTIONS", []):
        parts = line.split()
        if len(parts) >= 2:
            options[parts[0].upper()] = " ".join(parts[1:])
    start_date = options.get("START_DATE")
    start_time = options.get("START_TIME")
    end_date = options.get("END_DATE")
    end_time = options.get("END_TIME")
    if start_date:
        facts["sim_start"] = f"{start_date} {start_time}".strip() if start_time else start_date
    if end_date:
        facts["sim_end"] = f"{end_date} {end_time}".strip() if end_time else end_date

    # [SUBCATCHMENTS]: Name RainGage Outlet Area %Imperv Width Slope ...
    total_area = 0.0
    imperv_weighted = 0.0
    for line in sections.get("SUBCATCHMENTS", []):
        parts = line.split()
        if len(parts) < 5:
            continue
        try:
            area = float(parts[3])
            imperv = float(parts[4])
        except ValueError:
            continue
        total_area += area
        imperv_weighted += area * imperv
    if total_area > 0:
        facts["basin_area_ha"] = round(total_area, 2)
        facts["imperv_pct"] = round(imperv_weighted / total_area, 2)

    # [INFILTRATION] under Green-Ampt: Subcatch Suction Ksat IMD.
    # Only report these when the model actually uses Green-Ampt —
    # under Horton the same columns mean MaxRate/MinRate/Decay.
    infil_method = options.get("INFILTRATION", "").upper()
    if "GREEN" in infil_method:
        suctions: list[float] = []
        ksats: list[float] = []
        for line in sections.get("INFILTRATION", []):
            parts = line.split()
            if len(parts) < 3:
                continue
            try:
                suctions.append(float(parts[1]))
                ksats.append(float(parts[2]))
            except ValueError:
                continue
        if suctions:
            facts["suction_mm"] = round(sum(suctions) / len(suctions), 2)
        if ksats:
            facts["ksat_mm_per_hr"] = round(sum(ksats) / len(ksats), 2)

    return facts


def _load_artifacts(run_dir: str) -> dict:
    """Load all JSON artifacts from a run dir. Raises SystemExit on missing audit."""
    audit_dir = os.path.join(run_dir, "09_audit")
    if not os.path.isdir(audit_dir):
        print(
            f"ERROR: {audit_dir!r} not found. "
            "Run `aiswmm audit --run-dir <path>` before generating a report.",
            file=sys.stderr,
        )
        sys.exit(1)

    provenance_path = os.path.join(audit_dir, "experiment_provenance.json")
    if not os.path.exists(provenance_path):
        print(
            f"ERROR: {provenance_path!r} not found. "
            "The run must be fully audited before a report can be generated.",
            file=sys.stderr,
        )
        sys.exit(1)

    return {
        "run_dir": run_dir,
        "manifest": _load_json(os.path.join(run_dir, "manifest.json")),
        "inp_facts": _inp_model_facts(run_dir),
        "provenance": _load_json(provenance_path),
        "diagnostics": _load_json(os.path.join(run_dir, "model_diagnostics.json")),
        "comparison": _load_json(os.path.join(audit_dir, "comparison.json")),
        # Written by agentic_swmm.reporting.hydraulic_summary before this
        # script runs. This script is stdlib + python-docx + PyYAML only, so
        # the .rpt parsing happens on that side and arrives as JSON like every
        # other artifact here. Absent for runs generated before that step
        # existed; the section then says so instead of inventing numbers.
        "hydraulics": _load_json(os.path.join(audit_dir, "hydraulic_summary.json")),
        # Written by the design-review tool into the canonical 11_review
        # stage; absent when no review was run, and the section says so.
        "design_review": _load_json(os.path.join(run_dir, "11_review", "design_review.json")),
    }


# ---------------------------------------------------------------------------
# Main entry point
# ---------------------------------------------------------------------------

def generate(
    run_dir: str, out_path: str, template_path: str, title: str | None = None
) -> None:
    """Generate the Word deliverable. Pure function — no side-effects beyond writing out_path."""
    artifacts = _load_artifacts(run_dir)

    template = _load_template(template_path)
    sections = template.get("sections", [])
    if title:
        # CLI override beats the template's cover title so callers can
        # brand the deliverable per project without a custom template.
        for section_cfg in sections:
            if section_cfg.get("id") == "cover":
                section_cfg["title"] = title

    doc = Document()

    # Apply pure-black override to all Word built-in styles used by this generator
    _set_style_black(doc)

    # Insert page number field into the section footer
    _add_page_number_footer(doc)

    # Rendering context — mutable counters shared across all section renderers
    ctx = {
        "table_counter": [0],
        "figure_counter": [0],
        "section_number": 0,
    }

    for section_cfg in sections:
        section_id = section_cfg.get("id")
        if section_id not in SECTION_RENDERERS:
            print(
                f"WARNING: Unknown section id {section_id!r} in template — skipping.",
                file=sys.stderr,
            )
            continue
        SECTION_RENDERERS[section_id](doc, section_cfg, artifacts, ctx)

    # Ensure output directory exists
    out_dir = os.path.dirname(out_path)
    if out_dir:
        os.makedirs(out_dir, exist_ok=True)

    doc.save(out_path)


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Generate a Word deliverable from an audited SWMM run directory."
    )
    parser.add_argument("--run-dir", required=True, help="Path to the audited run directory.")
    parser.add_argument(
        "--out",
        default=None,
        help="Output .docx path (default: <run-dir>/report.docx).",
    )
    parser.add_argument(
        "--template",
        default=None,
        help="Path to a YAML (or JSON) template file. Defaults to the built-in default template.",
    )
    parser.add_argument(
        "--title",
        default=None,
        help="Override the cover title (default comes from the template).",
    )
    args = parser.parse_args()

    run_dir = os.path.abspath(args.run_dir)
    if not os.path.isdir(run_dir):
        print(f"ERROR: run-dir {run_dir!r} does not exist.", file=sys.stderr)
        sys.exit(1)

    out_path = args.out if args.out else os.path.join(run_dir, "report.docx")
    template_path = args.template if args.template else _default_template_path()

    generate(run_dir, out_path, template_path, title=args.title)
    print(f"Report written to: {out_path}")


if __name__ == "__main__":
    main()
